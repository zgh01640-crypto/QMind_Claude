"""
解析《房屋建筑与装饰工程工程量计算标准》DOCX 文件。

文档结构：
  正文段落（全为 Body Text / Normal 样式）
    附录 X  章名          → level=1 节点
    X.Y  节名             → level=2 节点
    表 X.Y.Z  表名(编码：XXXXXX)  → 接下来的表格属于此节
  表格（6列）：
    列0 项目编码（9位）  列1 项目名称  列2 项目特征
    列3 计量单位         列4 工程量计算规则  列5 工作内容

返回：
  sections: list[dict]   — {code, name, level, parent_code}
  items:    list[dict]   — {section_code, item_code, item_name, item_features,
                            unit, calc_rule, work_content}
"""
import re
from docx import Document


# ── 正则 ─────────────────────────────────────────────────────────────────────

_RE_APPENDIX = re.compile(r'^附录\s*([A-Z])\s+(.+)$')
_RE_SECTION  = re.compile(r'^([A-Z]\.\d+)\s+(.+)$')
_RE_TABLE_HDR = re.compile(r'表\s*([A-Z]\.\d+\.\d+).*?编码[：:]\s*(\d{6})')
_RE_ITEM_CODE = re.compile(r'^\d{9}$')


def _clean(text: str) -> str:
    return text.strip().replace('　', ' ')


def parse_docx(filepath: str) -> tuple[list[dict], list[dict]]:
    doc = Document(filepath)

    sections: list[dict] = []
    items: list[dict] = []

    # 将段落和表格按文档顺序排列
    # python-docx 的 doc.paragraphs / doc.tables 按顺序，但需要通过 body 获取混合顺序
    from docx.oxml.ns import qn
    body = doc.element.body

    current_appendix_code: str | None = None   # e.g. "A"
    current_section_code:  str | None = None   # e.g. "A.1"
    pending_section_code:  str | None = None   # 表格前最近一次 "表 X.Y.Z" 的 X.Y 部分
    section_num_code: dict[str, str] = {}      # alpha_code → 6位数字编码, e.g. "A.1" → "010101"

    # 附录→节 的父子关系在 sections 列表中按插入顺序体现
    # 用 dict 快速查父
    section_by_code: dict[str, str] = {}  # code -> parent_code

    para_idx = 0
    table_idx = 0

    def iter_body():
        """按 body 中的 XML 元素顺序，yield ('p', paragraph) 或 ('t', table)"""
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        p_nodes = doc.paragraphs
        t_nodes = doc.tables
        pi = ti = 0
        for child in body:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                if pi < len(p_nodes):
                    yield 'p', p_nodes[pi]
                    pi += 1
            elif tag == 'tbl':
                if ti < len(t_nodes):
                    yield 't', t_nodes[ti]
                    ti += 1
            elif tag == 'sdt':
                # 结构化文档标记，内部可能有段落/表格，跳过（通常是封面）
                pass

    for kind, node in iter_body():
        if kind == 'p':
            text = _clean(node.text)
            if not text:
                continue

            m = _RE_APPENDIX.match(text)
            if m:
                code = m.group(1)
                name = _clean(m.group(2))
                current_appendix_code = code
                current_section_code = None
                sections.append({'code': code, 'name': name, 'level': 1, 'parent_code': None})
                section_by_code[code] = None
                continue

            m = _RE_SECTION.match(text)
            if m:
                code = m.group(1)
                name = _clean(m.group(2))
                # 父是对应附录（code 前缀字母）
                parent = code.split('.')[0]
                current_section_code = code
                sections.append({'code': code, 'name': name, 'level': 2, 'parent_code': parent})
                section_by_code[code] = parent
                continue

            # "表 A.1.1 ... (编码：010101)" — 记下对应节编码
            m = _RE_TABLE_HDR.search(text)
            if m:
                tbl_ref = m.group(1)         # "A.1.1"
                num_code = m.group(2)        # "010101"
                # section code = 前两段, e.g. "A.1"
                parts = tbl_ref.split('.')
                pending_section_code = f"{parts[0]}.{parts[1]}" if len(parts) >= 2 else current_section_code
                if pending_section_code:
                    section_num_code[pending_section_code] = num_code
                continue

        else:  # kind == 't'
            table = node
            # 确定本表格所属节
            sec_code = pending_section_code or current_section_code

            rows = table.rows
            if len(rows) < 2:
                continue

            # 判断是否是清单表（首行/列头含"项目编码"）
            header_row = rows[0]
            header_cells = [_clean(c.text) for c in header_row.cells]
            is_header = any('项目编码' in hc or '项目名称' in hc for hc in header_cells)

            start_row = 1 if is_header else 0

            for row in rows[start_row:]:
                cells = [_clean(c.text) for c in row.cells]
                if len(cells) < 2:
                    continue
                code_cell = cells[0] if len(cells) > 0 else ''
                name_cell = cells[1] if len(cells) > 1 else ''

                # 跳过合并行（code列为空 or 不是9位编码）
                if not _RE_ITEM_CODE.match(code_cell.replace(' ', '')):
                    continue

                item_code = code_cell.replace(' ', '')
                item_name = name_cell

                item_features = cells[2] if len(cells) > 2 else ''
                unit          = cells[3] if len(cells) > 3 else ''
                calc_rule     = cells[4] if len(cells) > 4 else ''
                work_content  = cells[5] if len(cells) > 5 else ''

                # 去除重复空格/换行
                def norm(s: str) -> str:
                    return re.sub(r'\s+', ' ', s).strip()

                items.append({
                    'section_code':  sec_code,
                    'item_code':     item_code,
                    'item_name':     norm(item_name),
                    'item_features': norm(item_features) or None,
                    'unit':          norm(unit) or None,
                    'calc_rule':     norm(calc_rule) or None,
                    'work_content':  norm(work_content) or None,
                })

    for s in sections:
        s['num_code'] = section_num_code.get(s['code'])
    return sections, items


if __name__ == '__main__':
    import sys, json
    path = sys.argv[1] if len(sys.argv) > 1 else r'd:\ClaudeCode\QMind\mydoc\房屋建筑与装饰工程工程量计算标准.docx'
    secs, its = parse_docx(path)
    print(f'节: {len(secs)}  项目: {len(its)}')
    for s in secs[:5]:
        print(' ', s)
    for i in its[:5]:
        print(' ', i)
